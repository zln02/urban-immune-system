"""30인 자문 요청 요약표 워드 생성"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "30인_자문요청_요약표.docx")

EXPERTS = [
    ("01", "A. 감염병 역학/예방의학", "천병철 교수", "고려대학교 의과대학 예방의학교실", "통계 방법론 검증 — 17지역 walk-forward 다중검정 보정, Granger 검정 해석, Gate B 통계적 정당화"),
    ("02", "A. 감염병 역학/예방의학", "기모란 교수", "국립암센터 예방의학", "경보 임계값 및 보건당국 워크플로 적합성 — 실무 운영환경, YELLOW/RED 기준 현실성"),
    ("03", "A. 감염병 역학/예방의학", "조성일 교수", "서울대학교 의과대학 예방의학교실", "비의료 신호(OTC·검색)의 역학적 타당성 — 임상 확진 없는 조기경보 학술 정당화"),
    ("04", "A. 감염병 역학/예방의학", "최보율 교수", "한양대학교 의과대학 예방의학교실", "3계층 교차검증 설계 적절성 — walk-forward 5-fold 분할, 단일시즌 일반화 한계"),
    ("05", "A. 감염병 역학/예방의학", "나종화 교수", "충북대학교 의과대학 예방의학교실", "네이버 검색 트렌드 L3 신호 유효성 — 인플루엔자 감시 맥락에서 검색 신호 역학적 활용"),
    ("06", "A. 감염병 역학/예방의학", "정해관 교수", "성균관대학교 의과대학 예방의학교실/삼성서울병원", "17개 시·도 지역별 성능 차이 원인 — 빅데이터 의학연구 관점"),
    ("07", "A. 감염병 역학/예방의학", "이순영 교수", "충남대학교 의과대학 예방의학교실", "단일 시즌 데이터 일반화 한계 및 외부 검증 설계 권고"),
    ("08", "B. 하수역학/환경바이러스", "신용 교수", "연세대학교 공과대학 환경공학부", "KOWAS L2 신호 품질 — 충북·대구·경북 약신호 원인, WBE 분야 내 UIS 위치"),
    ("09", "B. 하수역학/환경바이러스", "유석현 교수", "가톨릭대학교 의과대학 (원헬스연구소)", "3계층 교차검증 원헬스 프레임 적합성 — 인수공통감염병 확장 가능성"),
    ("10", "B. 하수역학/환경바이러스", "강성원 박사", "한국환경연구원 (KEI) 물환경연구실", "하수처리장 운영 특성과 L2 신호 품질 관계 — 시설 규모·채취빈도 영향"),
    ("11", "B. 하수역학/환경바이러스", "오정익 교수", "한국기술교육대학교 에너지신소재화학공학부", "copies/mL 정규화 방법론 적절성 — 인구 보정 외 추가 보정 변수"),
    ("12", "C. AI/ML 헬스케어", "김재경 교수", "KAIST 수리과학과", "XGBoost 앙상블 + Gate B 조합 타당성 — 수리모델 기반 감염병 예측 관점"),
    ("13", "C. AI/ML 헬스케어", "박래웅 교수", "아주대학교 의과대학 의료정보학교실", "AI 경보 시스템 임상 의사결정 보조 적합성 — 의료정보학 관점"),
    ("14", "C. AI/ML 헬스케어", "황승식 교수", "서울대학교 보건대학원", "인터넷 기반 신호(L3 검색)의 역학적 활용 방법론 — 디지털 역학 관점"),
    ("15", "C. AI/ML 헬스케어", "이도헌 교수", "KAIST 바이오및뇌공학과", "TFT 시계열 예측 모델 및 Autoencoder 이상탐지 설계 — 바이오 AI 관점"),
    ("16", "C. AI/ML 헬스케어", "최윤섭 대표", "디지털헬스케어파트너스 (DHP)", "B2G 사업화 전략 및 의료 AI 규제 환경 — 헬스케어 스타트업 사업화"),
    ("17", "D. 수리생물학/감염병 모델링", "정은옥 교수", "건국대학교 수학과", "앙상블 가중치(w1=0.35, w2=0.40, w3=0.25) 수리적 근거 — SIR·SEIR 모델 관점"),
    ("18", "D. 수리생물학/감염병 모델링", "정일효 교수", "부산대학교 통계학과", "walk-forward 5-fold 결과 신뢰구간 및 가설검정 — 베이지안 추론 관점"),
    ("19", "D. 수리생물학/감염병 모델링", "이재용 교수", "서울대학교 통계학과", "17개 지역 다중비교(Multiple testing) 보정 필요성 — 통계적 추론"),
    ("20", "E. 보건정보학/공중보건 빅데이터", "이주연 교수", "연세대학교 보건대학원", "HIRA OpenAPI 연동 가능성 및 공공 보건 데이터 접근 방법"),
    ("21", "E. 보건정보학/공중보건 빅데이터", "박웅섭 교수", "아주대학교 보건대학원", "건강보험 청구 데이터를 L1 대체 신호로 활용 가능성"),
    ("22", "E. 보건정보학/공중보건 빅데이터", "임달오 교수", "경북대학교 보건학과", "약신호 지역(충북·대구·경북) 보건 인프라 특성 및 지역 임계값 보정 타당성"),
    ("23", "F. 공공기관/정부", "질병관리청 KOWAS팀", "질병관리청 감염병위기대응국 환경바이러스팀", "L2 데이터 품질·접근·인용 방식 + PoC 협력 가능성"),
    ("24", "F. 공공기관/정부", "국립감염병연구소 (NIID)", "질병관리청 산하", "ILINet 임상 감시 데이터와 UIS 경보 타이밍 상관 검증, 학술 접근 절차"),
    ("25", "F. 공공기관/정부", "한국보건산업진흥원 (KHIDI)", "보건복지부 산하", "감염병 AI 사업화 지원 프로그램 — 학생팀 신청 가능 여부, 기술가치 평가"),
    ("26", "F. 공공기관/정부", "보건복지부 디지털헬스과", "보건복지부", "B2G 조달 절차 및 스마트 방역 AI 공공 도입 정책"),
    ("27", "F. 공공기관/정부", "서울시 감염병관리지원단", "서울특별시 시민건강국", "광역시 단위 파일럿 PoC 협력 가능성"),
    ("28", "F. 공공기관/정부", "한국보건의료연구원 (NECA)", "보건복지부 산하", "AI 경보 시스템 의료기술 평가(HTA) 가이드라인 적합성"),
    ("29", "G. 국제기관/산업계", "WHO 서태평양지역사무처 (WPRO)", "세계보건기구 서태평양 지역", "WBE 기반 감염병 감시 국제 표준 및 글로벌 확장 가능성"),
    ("30", "G. 국제기관/산업계", "카카오헬스케어 AI연구소", "카카오헬스케어(주)", "민간-공공 연계 감염병 데이터 플랫폼 협력 — 데이터 연계·산학협력 가능성"),
]


def sf(run, size=10, bold=False):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._r.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")


doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

# 제목
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
r = title.add_run("UIS 전문가 자문 30인 요약표")
sf(r, size=16, bold=True)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(14)
r = subtitle.add_run("Urban Immune System Capstone 2026 — 각 전문가 채택 사유 및 자문 주제")
sf(r, size=10)

# 표
table = doc.add_table(rows=1, cols=5)
table.style = "Table Grid"

# 헤더
headers = ["No.", "분야", "전문가/기관", "소속", "자문 주제 (채택 사유)"]
widths = [Cm(1.0), Cm(3.2), Cm(3.2), Cm(4.5), Cm(7.5)]
for i, (cell, w) in enumerate(zip(table.rows[0].cells, widths)):
    cell.width = w
    p = cell.paragraphs[0]
    r = p.add_run(headers[i])
    sf(r, size=9, bold=True)
    p.paragraph_format.space_after = Pt(0)

# 데이터 행
current_group = ""
for no, group, name, org, topic in EXPERTS:
    # 분야 구분 행
    if group != current_group:
        current_group = group
        row = table.add_row()
        merged = row.cells[0].merge(row.cells[4])
        merged._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E1F2")
        shd.set(qn("w:val"), "clear")
        merged._tc.get_or_add_tcPr().append(shd)
        rg = merged.paragraphs[0].add_run(group)
        sf(rg, size=9, bold=True)
        merged.paragraphs[0].paragraph_format.space_after = Pt(0)

    row = table.add_row()
    data = [no, "", name, org, topic]
    for i, (cell, val) in enumerate(zip(row.cells, data)):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(val)
        sf(r, size=9)

doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
