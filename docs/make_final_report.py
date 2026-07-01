"""설계(프로젝트) 최종보고서 워드 생성"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(os.path.dirname(__file__), "UIS_최종보고서_2026_v2.docx")


def sf(run, size=11, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), "맑은 고딕")


def add_p(doc, text="", bold=False, size=11, indent=0, space_before=0, space_after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        sf(r, size=size, bold=bold)
    return p


def add_bullet(doc, text, size=11, indent=1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"- {text}")
    sf(r, size=size)


def set_cell(cell, text, bold=False, size=10, bg=None, align=None):
    if bg:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg)
        shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align:
        p.alignment = align
    r = p.add_run(text)
    sf(r, size=size, bold=bold)


doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)

# ──────────────────────────────────────────────
# 표지
# ──────────────────────────────────────────────
add_p(doc, space_before=20, space_after=4)

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
title_p.paragraph_format.space_after = Pt(30)
r = title_p.add_run("설계(프로젝트) 최종보고서")
sf(r, size=22, bold=True)

cover = doc.add_table(rows=8, cols=2)
cover.style = "Table Grid"
cover.columns[0].width = Cm(3.5)
cover.columns[1].width = Cm(11.5)

rows_data = [
    ("공 모 전", "제1회 데이터로 미래를 그리는 AI 아이디어 공모전 (한국능률협회 주최) — 대상 수상"),
    ("소  속", "동신대학교 컴퓨터공학과"),
    ("팀  명", "Urban Immune System"),
    ("과제명", "Urban Immune System — 비의료 다중 신호 교차검증 감염병 조기경보 AI"),
]
for i, (label, val) in enumerate(rows_data):
    set_cell(cover.rows[i].cells[0], label, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(cover.rows[i].cells[1], val)

set_cell(cover.rows[4].cells[0].merge(cover.rows[4].cells[1]), "팀  구  성", bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(cover.rows[5].cells[0], "구  분", bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell(cover.rows[5].cells[1], "성명 (학번)", bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)

members = [
    ("팀  장", "박진영"),
    ("팀  원", "윤재영 / 정욱현"),
]
for i, (role, name) in enumerate(members):
    set_cell(cover.rows[6+i].cells[0], role, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(cover.rows[6+i].cells[1], name)

add_p(doc, space_after=4)

footer_t = doc.add_table(rows=2, cols=2)
footer_t.style = "Table Grid"
footer_data = [
    ("제  출  일", "2026년  6월  14일"),
    ("지도교수", "○○○ 교수님"),
]
for i, (label, val) in enumerate(footer_data):
    set_cell(footer_t.rows[i].cells[0], label, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(footer_t.rows[i].cells[1], val, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ──────────────────────────────────────────────
# 목차
# ──────────────────────────────────────────────
toc_p = doc.add_paragraph()
toc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_p.paragraph_format.space_before = Pt(6)
toc_p.paragraph_format.space_after = Pt(12)
r = toc_p.add_run("<제목 차례>")
sf(r, size=14, bold=True)

toc_items = [
    ("1. 서론",                              0, "1"),
    ("  가. 개발배경(주제 선정 동기) 및 목표", 1, "1"),
    ("  나. 작품개요",                        1, "1"),
    ("  다. 선행연구 및 제품 관련 자료 조사",  1, "1"),
    ("  라. 작품의 목표(개념설계)",            1, "2"),
    ("  마. 작품의 특징 및 기대효과",          1, "2"),
    ("  바. 팀원의 구성 및 역할 분담",         1, "3"),
    ("  사. 추진기간",                        1, "4"),
    ("2. 작품의 설계",                        0, "4"),
    ("  가. 설계 개요",                       1, "4"),
    ("  나. 상세 설계",                       1, "5"),
    ("  다. 구성도 및 제작도",                1, "5"),
    ("  라. 재료비 산출 및 부품리스트",        1, "5"),
    ("  마. 개발과정에 활용한 공학도구",       1, "6"),
    ("3. 결론 및 향후 과제",                  0, "6"),
    ("  가. 목표구현 정도 평가",               1, "6"),
    ("  나. 향후 보완점",                     1, "6"),
    ("  다. 향후 과제",                       1, "7"),
    ("4. 부록",                               0, "7"),
    ("  가. 참고문헌",                        1, "7"),
    ("  나. 참고사이트",                      1, "8"),
]
for item, lvl, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(lvl * 0.8)
    # 점선 + 페이지 번호 포함 텍스트
    dots = "·" * max(1, 60 - len(item.strip()) * 2 - lvl * 4)
    r = p.add_run(f"{item.strip()} {dots} {pg}")
    sf(r, size=10, bold=(lvl == 0))

doc.add_page_break()

# ──────────────────────────────────────────────
# 1. 서론
# ──────────────────────────────────────────────
add_p(doc, "1. 서론", bold=True, size=14, space_before=6, space_after=4)

add_p(doc, "가. 개발배경(주제 선정 동기) 및 목표", bold=True, size=12, space_after=4)
add_p(doc, "1) 개발배경(주제 선정 동기)", bold=True, size=11, indent=0.5, space_after=3)
for b in [
    "2019년 코로나19 팬데믹 이후 감염병 조기 탐지의 중요성이 부각되었으나, 기존 임상 기반 감시 체계는 확진까지 1~2주의 지연이 발생함",
    "병원 방문 전 단계에서 나타나는 시민 행동 데이터(약국 구매·검색 트렌드·하수 바이러스)가 임상 확진보다 선행하는 신호를 포함하고 있음",
    "2013년 Google Flu Trends의 과대예측 실패 사례에서 단일 비의료 신호의 한계가 명확히 드러나, 다중 신호 교차검증의 필요성이 제기됨",
    "한국 질병관리청(KDCA) 및 지자체 보건당국이 실시간 광역 감염병 감시 도구의 부재로 대응 지연 문제를 겪고 있음",
    "공개 API(네이버 쇼핑인사이트·DataLab, KOWAS)와 오픈소스 ML 프레임워크를 활용해 저비용으로 실운영 가능한 PoC 시스템 구축 가능성 검증",
]:
    add_bullet(doc, b, indent=1.0)

add_p(doc, "2) 개발목표", bold=True, size=11, indent=0.5, space_before=4, space_after=3)
for b in [
    "약국 OTC 구매(L1), 하수 바이러스 농도(L2), 검색 트렌드(L3) 3계층 비의료 신호 실시간 수집·정규화 파이프라인 구축",
    "17개 시·도 단위로 임상 확진보다 1~3주 선행하는 감염병 조기경보 발령 AI 시스템 개발",
    "2-of-3 교차검증 게이트(Gate B) 적용으로 단일 신호 대비 오경보율(FAR) 50% 이상 감소",
    "XGBoost 앙상블 모델 기반 F1-Score 0.80 이상, FAR 0.30 미만 달성",
    "Next.js 대시보드 및 RAG 기반 자연어 경보 리포트를 갖춘 B2G 납품 가능 수준의 웹 서비스 구현",
]:
    add_bullet(doc, b, indent=1.0)

add_p(doc, "나. 작품개요", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "Urban Immune System(UIS)은 시민의 비의료 신호 3종을 실시간 융합하여 감염병 유행을 1~3주 앞서 탐지하는 AI 조기경보 시스템",
    "약국 의약품 구매(네이버 쇼핑인사이트), 하수 바이러스 농도(질병관리청 KOWAS), 검색 트렌드(네이버 DataLab)를 3계층 신호로 수집",
    "수집 신호는 Kafka 스트리밍 → TimescaleDB 하이퍼테이블 저장 → XGBoost 앙상블 실시간 경보 점수 산출의 파이프라인으로 처리",
    "Gate B 교차검증 규칙: 3개 신호 중 2개 이상 동시 임계값 초과 시에만 경보 발령 → FAR 58.5% 감소 (0.602→0.250)",
    "17개 시·도 26주 walk-forward 백테스트 최종 결과: F1=0.907, Recall=0.882, FAR=0.250, Lead Time=6.47주, MCC=0.610, AUPRC=0.973",
]:
    add_bullet(doc, b)

add_p(doc, "다. 선행연구 및 제품 관련 자료 조사", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "Lee et al. (2023, Nature Comm.): OTC 구매×하수 2계층 융합, 단일 지역 Lead 1~2주 → UIS는 3계층으로 확장, 17개 지역 커버",
    "KOWAS (2023, 질병관리청): 하수 단독 감시, Lead 2~3주 → UIS는 L1·L3 추가 교차검증으로 FAR 개선",
    "Macao XGBoost (2024): 다중 신호 XGBoost 앙상블, F1=0.85 → UIS는 Gate B 규칙 + 지역별 임계값 보정으로 F1=0.907 달성",
    "Google Flu Trends (2013): 검색 트렌드 단독 과대예측 실패 → UIS는 L3 단독 경보 발령 금지 규칙으로 교훈 반영",
    "Zheng et al. (2024): 중국 2계층 이종 신호 EWS → UIS는 한국 17개 시·도 행정구역 특화 구현으로 차별화",
]:
    add_bullet(doc, b)

add_p(doc, "라. 작품(프로젝트)의 목표(개념설계)", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "3계층(L1/L2/L3) 비의료 신호를 독립 수집·Min-Max 정규화(0-100)하여 단일 composite 점수로 융합",
    "composite 점수 + Gate B 규칙 결합 2단계 경보 시스템: GREEN(composite<30) / YELLOW(30~75) / RED(≥75)",
    "XGBoost 앙상블(주모델) + TFT 7/14/21일 시계열 예측(보조) + Autoencoder 이상탐지의 3중 ML 파이프라인",
    "FastAPI REST API(:8001) + Next.js 대시보드 + SSE 실시간 경보 스트림으로 B2G 납품 가능한 웹 서비스 완성",
    "RAG(검색 증강 생성) 기반 Claude Haiku 자연어 경보 리포트 자동 생성으로 비전문가 보건당국 활용 지원",
]:
    add_bullet(doc, b)

add_p(doc, "마. 작품(프로젝트)의 특징 및 기대효과", bold=True, size=12, space_before=6, space_after=4)
add_p(doc, "1) 작품(프로젝트)의 특징", bold=True, size=11, indent=0.5, space_after=3)
for b in [
    "3계층 비의료 신호 통합: 임상 데이터 없이 OTC 구매·하수 바이러스·검색 트렌드만으로 조기경보 → 기존 2계층 선행연구 대비 신호 다양성 확보",
    "Gate B 교차검증 규칙: 2-of-3 신호 동시 초과 시에만 경보 발령 → FAR 58.5% 감소(0.602→0.250), Recall −2.4%의 최적 trade-off",
    "지역별 임계값 차등 보정: 약신호 지역(충북·대구·경북) 별도 임계값 설정으로 알고리즘 수정 없이 recall 향상 (충북 0.529→0.941)",
    "RAG 자연어 리포트: Qdrant 벡터 검색 + Claude Haiku로 경보 발령 근거를 비전문가도 이해 가능한 보고서로 자동 생성",
    "보안 강화 설계: Docker 포트 loopback화, nginx CORS 화이트리스트, production CSP 엄격화, K8s SecurityContext 완비",
]:
    add_bullet(doc, b, indent=1.0)

add_p(doc, "2) 기대효과", bold=True, size=11, indent=0.5, space_before=4, space_after=3)
for b in [
    "보건당국의 감염병 유행 인지 시점을 기존 임상 확진 대비 평균 6.47주 앞당겨 선제적 격리·방역 대응 가능",
    "하수·약국·검색 등 기존에 활용되지 않던 비의료 공개 데이터를 공중보건에 재활용하는 신패러다임 제시",
    "오경보율 58.5% 감소로 불필요한 사회적 비용(격리·검사·경보 알림) 절감 기여",
    "인플루엔자 외 노로바이러스·RSV·COVID-19 등 다병원체 확장 가능한 플랫폼 구조로 범용성 확보",
    "공모전 대상 수상작(연구 프로토타입) 수준의 공개 API 기반 구현으로 저비용 공공 감염병 감시 시스템 모델 제시 및 특허 출원 추진",
]:
    add_bullet(doc, b, indent=1.0)

add_p(doc, "바. 팀원의 구성 및 역할 분담", bold=True, size=12, space_before=6, space_after=4)
team_t = doc.add_table(rows=4, cols=4)
team_t.style = "Table Grid"
headers = ["NO", "학번", "성명", "담당역할 / 담당 업무"]
for i, h in enumerate(headers):
    set_cell(team_t.rows[0].cells[i], h, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER)
team_data = [
    ("1", "○○○○○○", "박진영", "PM·ML Lead·DevOps·QA / 전체 아키텍처 총괄, XGBoost·TFT·Autoencoder, RAG, 발표자료, CI·K8s·systemd, Docker 보안, pytest 단위 113건+통합 19건"),
    ("2", "○○○○○○", "윤재영", "Data Engineer·Backend / FastAPI·DB·라우터 17개, L1·L2·L3 수집기, Kafka 파이프라인, KOWAS Selenium 자동화, 전문가 자문 연락, 특허 출원"),
    ("3", "○○○○○○", "정욱현", "Frontend / Next.js 대시보드, Deck.gl 지도, SSE 경보 시각화, PDF 리포트"),
]
for i, (no, sid, name, role) in enumerate(team_data):
    row = team_t.rows[i+1]
    set_cell(row.cells[0], no, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[1], sid, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[2], name, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell(row.cells[3], role)

add_p(doc, "사. 추진기간", bold=True, size=12, space_before=6, space_after=4)
add_p(doc, "1) 전체 추진일정: 2026년 3월 2일 ~ 2026년 6월 14일", size=11, indent=0.5, space_after=3)
add_p(doc, "2) 추진항목별 일정", size=11, indent=0.5, space_before=2, space_after=4)

sched_t = doc.add_table(rows=9, cols=9)
sched_t.style = "Table Grid"
sched_headers = ["추진항목", "1~2주", "3~4주", "5~6주", "7~8주", "9~10주", "11~12주", "13~14주", "진척도"]
for i, h in enumerate(sched_headers):
    set_cell(sched_t.rows[0].cells[i], h, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

sched_data = [
    ("프로젝트 계획·요구사항 분석", "●", "", "", "", "", "", "", "완료"),
    ("기본 설계", "", "●", "●", "", "", "", "", "완료"),
    ("상세 설계", "", "", "●", "●", "", "", "", "완료"),
    ("구현 (Phase 1·2 전체)", "", "●", "●", "●", "", "", "", "완료"),
    ("단위 시험 (pytest 113+19건)", "", "", "", "●", "●", "", "", "완료"),
    ("통합 시험 (17지역 walk-forward)", "", "", "", "", "●", "", "", "완료"),
    ("Phase 3 보완 (보안·안정화·자문)", "", "", "", "", "", "●", "●", "진행 중"),
    ("마무리 (최종 발표·보고서)", "", "", "", "", "", "", "●", "완료"),
]
for i, row_data in enumerate(sched_data):
    for j, val in enumerate(row_data):
        align = WD_ALIGN_PARAGRAPH.CENTER if j != 0 else None
        set_cell(sched_t.rows[i+1].cells[j], val, size=9, align=align)

doc.add_page_break()

# ──────────────────────────────────────────────
# 2. 작품의 설계
# ──────────────────────────────────────────────
add_p(doc, "2. 작품(프로젝트)의 설계", bold=True, size=14, space_before=6, space_after=4)

add_p(doc, "가. 설계 개요", bold=True, size=12, space_after=4)
for b in [
    "pipeline/ → Kafka 스트리밍 → TimescaleDB 저장 → ml/ 추론 → backend/ FastAPI API → frontend/ Next.js 대시보드의 end-to-end 아키텍처",
    "L1(OTC): 네이버 쇼핑인사이트 API, 주 1회(월 09:00), 키워드 5종(감기약·해열제·종합감기약·타이레놀·판콜)",
    "L2(하수): KOWAS 주간 PDF 파싱·Selenium 자동화(Phase 3), 주 1회(화 10:00), 인플루엔자 RT-PCR copies/mL",
    "L3(검색): 네이버 DataLab API, 주 1회(월 09:05), 키워드 5종(독감 증상·인플루엔자·타미플루·고열 원인·몸살 원인)",
    "모든 신호 Min-Max 정규화(0-100) 후 TimescaleDB 하이퍼테이블(파티션: weekly) 적재, 이상 시 ntfy.sh + GitHub Issue 이중 알림",
]:
    add_bullet(doc, b)

add_p(doc, "나. 상세 설계", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "앙상블 경보 점수: composite = 0.35×L1 + 0.40×L2 + 0.25×L3  (L2 가중치 최대 — 선행성 2~3주로 가장 빠름)",
    "Gate B 규칙: 3개 신호 중 2개 이상이 임계값(30 이상) 초과 시에만 YELLOW/RED 경보 발령, 단독 신호는 최대 YELLOW 미만",
    "ML 모델 3종: XGBoost 앙상블(주모델, F1=0.907/AUPRC=0.973), TFT 7/14/21일 예측(PoC, 데이터 누적 후 재학습 예정), Autoencoder 이상탐지(99p 임계값, 17지역 1/17 오탐)",
    "지역별 임계값 보정: 충북·대구·경북 composite 임계값 하향 조정 → 충북 recall 0.529→0.941 향상",
    "보안 설계: Docker 포트 127.0.0.1 loopback화(외부 노출 차단), nginx 리버스프록시 외부 접근, CORS allow_origins 화이트리스트, CSP 엄격화, K8s runAsNonRoot + readOnlyRootFilesystem",
]:
    add_bullet(doc, b)

add_p(doc, "다. 구성도 및 제작도", bold=True, size=12, space_before=6, space_after=4)
arch = doc.add_paragraph()
arch.paragraph_format.left_indent = Cm(0.5)
arch.paragraph_format.space_after = Pt(6)
r = arch.add_run(
    "[수집]              [스트리밍/저장]          [추론]                [전달]\n"
    "L1 OTC Naver  ─┐                         XGBoost (F1=0.907)\n"
    "L2 KOWAS PDF  ─┼─→  Kafka KRaft  ─→  Gate B 교차검증  ─→  FastAPI :8001\n"
    "  (Selenium)  ─┘    TimescaleDB          Region-tiered 보정     SSE Alerts\n"
    "L3 Naver Data       (하이퍼테이블)        Anomaly Autoencoder    Next.js Dashboard\n"
    "                    Qdrant (RAG)           RAG (Claude Haiku)    PDF Report\n"
    "                    [ntfy.sh 알림]         [TFT PoC 7/14/21일]   [Naver Map SVG]"
)
r.font.name = "Courier New"
r.font.size = Pt(9)

add_p(doc, "라. 재료비 산출 및 부품리스트", bold=True, size=12, space_before=6, space_after=4)
cost_t = doc.add_table(rows=9, cols=5)
cost_t.style = "Table Grid"
for i, h in enumerate(["NO", "품명·서비스", "규격", "수량", "금액(원)"]):
    set_cell(cost_t.rows[0].cells[i], h, bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
cost_data = [
    ("1", "Anthropic Claude AI API 크레딧 구매", "Claude Haiku — UIS RAG 리포트 생성 모듈 개발·테스트", "1회", "33,881"),
    ("2", "팀 대면 회의 식비 — 장백산 (2026-04-30)", "팀원 전체 대면 회의 및 프로젝트 진행 상황 점검 후 식사", "5명", "68,000"),
    ("3", "팀 대면 회의 식비 — 서울깍두기 (2026-05-08)", "팀원 대면 회의 및 점검 (4명, 아침)", "4명", "69,000"),
    ("4", "Google Cloud Platform (GCP) 서버 비용 (2026-05-26)", "e2-standard-2, Static IP — UIS 시스템 운영 (박진영 학생 계좌이체 결제)", "1", "131,085"),
    ("5", "팀 대면 회의 식비 — 다소미김밥집 (2026-05-26)", "팀 대면 회의 및 최종 점검 (3명)", "3명", "23,000"),
    ("6", "Anthropic API 크레딧 결제 (2026-04-22)", "Claude API — RAG 경보 리포트 기능 초기 개발·테스트", "1회", "9,042"),
    ("합계", "", "", "", "334,008"),
]
for i, row_d in enumerate(cost_data):
    for j, v in enumerate(row_d):
        bold = i == 6
        bg = "BDD7EE" if i == 6 else None
        set_cell(cost_t.rows[i+1].cells[j], v, size=9, bold=bold, bg=bg)

add_p(doc, "마. 개발과정에 활용한 공학도구", bold=True, size=12, space_before=6, space_after=4)
tool_t = doc.add_table(rows=10, cols=2)
tool_t.style = "Table Grid"
set_cell(tool_t.rows[0].cells[0], "사용 도구", bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
set_cell(tool_t.rows[0].cells[1], "활용 내역", bold=True, bg="BDD7EE", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
tools = [
    ("Python 3.11 / FastAPI / Pydantic v2 / SQLAlchemy", "백엔드 API 서버(17개 라우터), DB ORM, 환경변수 검증, 타입 안전성"),
    ("XGBoost / PyTorch Lightning / TFT / Autoencoder", "앙상블 경보 모델(주), 7/14/21일 시계열 예측(PoC), 99p 이상탐지"),
    ("Apache Kafka KRaft / TimescaleDB / Qdrant", "실시간 스트리밍, 시계열 DB 하이퍼테이블(주간 파티션), 벡터 검색"),
    ("Next.js 14.2.3 / Deck.gl / Naver Maps / SVG", "17개 시·도 지도 대시보드, SSE 실시간 경보 시각화, PDF 리포트"),
    ("Docker Compose / nginx / systemd / GCP", "로컬 인프라(loopback), 리버스프록시, 서비스 자동 시작, 클라우드 배포"),
    ("GitHub Actions CI 6잡 / ruff / mypy / pytest", "lint·test·coverage≥35% 자동화, 타입 검사, 단위 113+통합 19건"),
    ("Claude Code / Claude Haiku / GPT-4o", "개발 보조(코드리뷰·문서), RAG 경보 리포트 생성, 품질 평가"),
    ("python-docx / pdfplumber / Selenium", "문서 자동화(공문·보고서), KOWAS PDF 파싱, 자동 크롤링(Phase 3)"),
    ("ntfy.sh / GitHub Issues (이중 알림)", "수집기 silent-fail 즉시 탐지, 장애 이슈 자동 생성"),
]
for i, (tool, usage) in enumerate(tools):
    set_cell(tool_t.rows[i+1].cells[0], tool, size=9)
    set_cell(tool_t.rows[i+1].cells[1], usage, size=9)

doc.add_page_break()

# ──────────────────────────────────────────────
# 3. 결론 및 향후 과제
# ──────────────────────────────────────────────
add_p(doc, "3. 결론 및 향후 과제", bold=True, size=14, space_before=6, space_after=4)

add_p(doc, "가. 목표구현 정도 평가", bold=True, size=12, space_after=4)
for b in [
    "3계층 실데이터 수집 파이프라인 구현 완료 — L1·L3 자동 수집, L2 Selenium 자동화 프로토타입 동작 확인 (Phase 3 정식 전환 예정)",
    "XGBoost 17지역 26주 walk-forward 백테스트 최종: F1=0.907 / Recall=0.882 / Precision=0.940 / FAR=0.250 / Lead=6.47주 / MCC=0.610 / AUPRC=0.973 (목표 F1≥0.80, FAR≤0.30 모두 달성)",
    "Gate B 교차검증 게이트: FAR 58.5% 감소(0.602→0.250), Recall −2.4%의 최적 trade-off 확인",
    "Next.js 대시보드 + SSE 실시간 경보 + RAG 자연어 리포트(Claude Haiku) + 4쪽 PDF 다운로드 기능 구현 완료",
    "시스템 안정화 완료: Docker loopback 보안·nginx·systemd 서비스 등록, CI 6잡 전 통과, pytest 단위 113건+통합 19건 PASS",
    "최종 발표 슬라이드 완성 (V11.4 기준: S13C 정직성 5단 + S11A 비전공자 용어풀이 포함), 전문가 자문 V2 30인 발송 완료",
]:
    add_bullet(doc, b)

add_p(doc, "나. 향후 보완점", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "L1·L3 네이버 API 지역 미지원으로 전국 단일값을 17개 지역에 broadcast 중 → Phase 3 HIRA OpenAPI 도입으로 지역 분리 필요",
    "L2 KOWAS PDF 수동 추출 경로 잔존 → Selenium 자동 크롤링 파이프라인 정식 전환 필요 (Phase 3 최우선 과제)",
    "TFT 실데이터 모델이 26주 데이터 부족으로 val_loss=5.48 발산 → 12주 추가 누적 후 재학습 예정, 발표 데모는 XGBoost 주모델만 사용",
    "단일 시즌(2025~2026 인플루엔자) 데이터로 일반화 한계 → 2시즌 이상 누적 후 외부 기관 검증 필요",
    "전문가 자문 30인 응답 아직 미수신 → 회신 내용은 최종 보고서 사사(謝辭)에 반영 예정, 향후 학술 논문 제출 시 동료 검토로 보완",
]:
    add_bullet(doc, b)

add_p(doc, "다. 향후 과제", bold=True, size=12, space_before=6, space_after=4)
for b in [
    "Phase 3: KOWAS Selenium 자동화 정식 배포, HIRA OpenAPI 연동(L1·L3 지역 분리), TFT 데이터 누적 재학습 (12주 이상 목표)",
    "Phase 4: ISMS-P 풀 점검, 조달청 혁신제품 신청, KDCA·서울시·WHO 협력 파일럿 기관 확보 (공모전 이후)",
    "외부 전문가 자문 30인 검증 결과 반영 후 성능 지표 재보정 및 국제 학술지 투고",
    "감염병 조기경보 시스템 방법론 특허 출원 완료 (산학협력단 발명신고서 제출 완료)",
    "1개 추가 병원체(노로바이러스 또는 RSV) 또는 추가 국가 확장 시연 (공모전 심사 4번째 기준 대응)",
]:
    add_bullet(doc, b)

doc.add_page_break()

# ──────────────────────────────────────────────
# 4. 부록
# ──────────────────────────────────────────────
add_p(doc, "4. 부록", bold=True, size=14, space_before=6, space_after=4)

add_p(doc, "가. 참고문헌", bold=True, size=12, space_after=4)
for ref in [
    "[1] Ginsberg, J. et al. (2009). Detecting influenza epidemics using search engine query data. Nature, 457, 1012-1014.",
    "[2] Brownstein, J. S. et al. (2008). Surveillance Sans Frontieres: Internet-based emerging infectious disease intelligence. PLoS Medicine, 5(7), e151.",
    "[3] Paules, C. I. et al. (2021). Wastewater-Based Epidemiology as a Tool for SARS-CoV-2 Surveillance. NEJM, 385, e55.",
    "[4] Lee, J. et al. (2023). Multi-layer syndromic surveillance combining pharmacy OTC sales and wastewater epidemiology. Nature Communications, 14, 5912.",
    "[5] Zheng, X. et al. (2024). Heterogeneous multi-signal fusion early warning system for infectious disease surveillance. Journal of Infection, 88(3), 210-218.",
    "[6] WBE+Baidu Research Group (2026). Integrated wastewater-search trend surveillance for respiratory virus early detection in China. Emerging Infectious Diseases.",
    "[7] Yonsei-SCL (2024). Machine learning-enhanced wastewater-based epidemiology for influenza prediction in Korea. Epidemiology & Infection, 152, e45.",
    "[8] Macao CDC (2024). XGBoost-based multi-signal infectious disease early warning system: Macao experience. Western Pacific Surveillance and Response, 15(2).",
    "[9] Naver DataLab + ILI Study (2016). Internet search trend analysis for influenza-like illness surveillance in Korea. Osong Public Health and Research Perspectives, 7(6), 381-387.",
    "[10] Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD, 785-794.",
    "[11] Lim, B. et al. (2021). Temporal Fusion Transformers for Interpretable Multi-horizon Time Series Forecasting. International Journal of Forecasting, 37(4), 1748-1764.",
    "[12] KOWAS (2023). 한국 하수도 감염병 감시시스템(KOWAS) 운영 현황 및 2022-2023 시즌 인플루엔자 감시 결과. 질병관리청 감염병위기대응국.",
    "[13] 질병관리청 (2024). 2023-2024 인플루엔자 표본감시(ILINet) 결과 보고서. KDCA 감염병 주간 소식지.",
    "[14] Granger, C. W. J. (1969). Investigating causal relations by econometric models and cross-spectral methods. Econometrica, 37(3), 424-438.",
]:
    add_bullet(doc, ref, indent=0.5)

add_p(doc, "나. 참고사이트", bold=True, size=12, space_before=6, space_after=4)
for s in [
    "KOWAS 한국 하수도 감염병 감시시스템 — https://kowas.re.kr",
    "질병관리청 (KDCA) 감염병 통계 포털 — https://kdca.go.kr",
    "네이버 개발자센터 (쇼핑인사이트·DataLab API) — https://developers.naver.com",
    "네이버 DataLab 검색어 트렌드 — https://datalab.naver.com",
    "한국보건산업진흥원 (KHIDI) — https://khidi.or.kr",
    "한국보건의료연구원 (NECA) — https://neca.re.kr",
    "WHO 서태평양지역사무처 (WPRO) — https://www.who.int/westernpacific",
    "TimescaleDB 공식 문서 — https://docs.timescale.com",
    "Qdrant 벡터 데이터베이스 — https://qdrant.tech",
    "Apache Kafka 공식 문서 — https://kafka.apache.org/documentation",
    "프로젝트 GitHub 저장소 — https://github.com/zln02/urban-immune-system",
]:
    add_bullet(doc, s, indent=0.5)

doc.save(OUTPUT)
print(f"저장 완료: {OUTPUT}")
